import os
import fitz  # PyMuPDF
from docx import Document
from typing import Tuple, Optional, List, Dict, Any
from pathlib import Path
import logging
from text_analyzer import TextAnalyzer

# Cấu hình logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """Xử lý các loại file khác nhau để trích xuất văn bản cho hệ thống UPT"""
    
    # [CẬP NHẬT] Mở rộng danh sách file hỗ trợ (Code & Văn bản)
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.doc': 'Word Document',
        '.txt': 'Text File',
        '.py': 'Python Code',
        '.js': 'JavaScript Code',
        '.html': 'HTML File',
        '.css': 'CSS File',
        '.json': 'JSON Data',
        '.md': 'Markdown File',
        '.csv': 'CSV Data'
    }

    # Cache cho TextAnalyzer để tránh khởi tạo nhiều lần
    _analyzer_instance = None

    @classmethod
    def _get_analyzer(cls):
        """Lazy loading cho TextAnalyzer"""
        if cls._analyzer_instance is None:
            logger.info("Đang khởi tạo TextAnalyzer...")
            cls._analyzer_instance = TextAnalyzer(language='vietnamese')
        return cls._analyzer_instance
    
    @classmethod
    def is_supported_file(cls, file_path: str) -> bool:
        """Kiểm tra xem file có được hỗ trợ không"""
        ext = Path(file_path).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def get_file_type(file_path: str) -> str:
        """Xác định loại file dựa trên phần mở rộng"""
        return Path(file_path).suffix.lower()
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Trích xuất văn bản từ file PDF sử dụng PyMuPDF"""
        text = ""
        doc = None
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text("text") + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Lỗi khi đọc PDF {file_path}: {str(e)}")
            return ""
        finally:
            if doc:
                doc.close()
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Trích xuất văn bản từ file DOCX (Bao gồm cả Paragraphs và Tables)"""
        text_content = []
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            if doc.tables:
                text_content.append("\n[NỘI DUNG TỪ BẢNG BIỂU]:")
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_content.append(" | ".join(row_text))
                    text_content.append("-" * 20) 

            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Lỗi khi đọc file DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Đọc nội dung file TXT và Code (xử lý nhiều encoding)"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Lỗi khi đọc file TXT/Code: {str(e)}")
                return ""
        return ""
    
    @classmethod
    def process_file(cls, file_path: str, analyze: bool = True) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
        """Xử lý file và trích xuất văn bản"""
        if not os.path.exists(file_path):
            return False, f"File không tồn tại: {file_path}", None
            
        try:
            file_ext = cls.get_file_type(file_path)
            logger.info(f"Đang xử lý file: {file_path} [{file_ext}]")
            
            if file_ext not in cls.SUPPORTED_EXTENSIONS:
                return False, f"Định dạng {file_ext} chưa được hỗ trợ.", None
            
            # Xử lý trích xuất
            text = ""
            if file_ext == '.pdf':
                text = cls.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                if file_ext == '.doc':
                    return False, "Định dạng .doc cũ không được hỗ trợ, vui lòng đổi sang .docx hoặc .pdf", None
                text = cls.extract_text_from_docx(file_path)
            # [CẬP NHẬT] Xử lý tất cả các file code như file text
            elif file_ext in ['.txt', '.py', '.js', '.html', '.css', '.json', '.md', '.csv']:
                text = cls.extract_text_from_txt(file_path)
            
            if not text or not text.strip():
                return False, "Không trích xuất được nội dung (File rỗng hoặc lỗi)", None
            
            # Chuẩn bị kết quả
            result = {
                'content': text.strip(),
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'file_type': file_ext
            }
            
            # Phân tích nội dung (chỉ phân tích nếu file không phải là code thuần túy để tránh nhiễu, hoặc tùy chọn)
            # Ở đây ta vẫn cho phân tích để Deloris biết đây là file "Code/Công nghệ"
            if analyze and len(text.strip()) > 50: 
                try:
                    analyzer = cls._get_analyzer()
                    analysis = analyzer.analyze(text)
                    
                    result.update({
                        'analysis': {
                            'summary': analysis.get('summary', 'N/A'),
                            'keywords': [{'word': k[0], 'count': k[1]} for k in analysis.get('keywords', [])],
                            'category': analysis.get('category', 'Uncategorized')
                        }
                    })
                except Exception as e:
                    logger.error(f"Lỗi khi phân tích văn bản: {str(e)}", exc_info=True)
                    result['analysis'] = {'error': 'Lỗi phân tích', 'details': str(e)}
            
            return True, "Xử lý thành công", result
            
        except Exception as e:
            error_msg = f"Lỗi nghiêm trọng khi xử lý file: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return False, error_msg, None
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        return list(cls.SUPPORTED_EXTENSIONS.keys())